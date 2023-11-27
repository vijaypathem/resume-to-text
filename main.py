from fastapi import FastAPI, File, UploadFile, HTTPException, Request
from fastapi.responses import HTMLResponse
import fitz
import csv
from bs4 import BeautifulSoup
from typing import List
from fastapi.templating import Jinja2Templates
from docx import Document
import io
import docx


app = FastAPI()
templates = Jinja2Templates(directory="templates")

def extract_and_clean_text(file_contents: bytes, file_extension: str) -> str:
    text = ''
    if file_extension == 'pdf':
        with fitz.open(stream=file_contents, filetype="pdf") as doc:
            for page in doc:
                text += page.get_text()
    elif file_extension == 'html':
        soup = BeautifulSoup(file_contents, 'html.parser')
        text = soup.get_text()
    elif file_extension == 'docx':
        doc = docx.Document(io.BytesIO(file_contents))
        for paragraph in doc.paragraphs:
            text += paragraph.text + " "
        for header in doc.sections:
            for paragraph in header.header.paragraphs:
                text += paragraph.text + " "
        for footer in doc.sections:
            for paragraph in footer.footer.paragraphs:
                text += paragraph.text + " "

    cleaned_text = ' '.join(text.split())
    return cleaned_text

def append_to_csv(cleaned_text: str):
    with open('dataset/resumedata.csv', 'a', newline='', encoding='utf-8') as csv_file:
        csv_writer = csv.writer(csv_file)
        if csv_file.tell() == 0:
            csv_writer.writerow(['ResumeData'])
        csv_writer.writerow([cleaned_text])

@app.post("/uploadfile/")
async def create_upload_file(request: Request, files: List[UploadFile] = File(...)):
    if len(files) == 0:
        message = "Please upload a file."
        return templates.TemplateResponse("upload_form.html", {"request": request, "message": message})

    message = ""
    for uploaded_file in files:
        file_contents = await uploaded_file.read()
        file_extension = uploaded_file.filename.split('.')[-1].lower()
        cleaned_text = extract_and_clean_text(file_contents, file_extension)
        append_to_csv(cleaned_text)
        message = "success"

    return templates.TemplateResponse("upload-form.html", {"request": request, "message": message})

@app.get("/", response_class=HTMLResponse)
def read_form(request: Request):
    return templates.TemplateResponse("upload-form.html", {"request": request, "message": ""})
